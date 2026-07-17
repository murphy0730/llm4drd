# -*- coding: utf-8 -*-
"""把结构化 BLOCKS 组装成 Word 文档。

BLOCK 格式（chapters/ch*.py 中定义 BLOCKS 列表）：
  ("h1"|"h2"|"h3", 文本)
  ("p", 文本)            支持 **加粗** 与 `行内代码`
  ("quote", 文本)        关键定义，引用样式
  ("code", 文本)         代码块/ASCII 图（等宽、灰底、保留换行）
  ("formula", 文本)      公式，居中等宽
  ("table", {"headers": [...], "rows": [[...]], "caption": 可选})
  ("bullets", [文本...]) 无序列表
  ("numbers", [文本...]) 有序列表
  ("pagebreak",)
"""
from __future__ import annotations

import datetime
import importlib
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOC_TITLE = "LLM4DRD 排产算法基础理论"
DOC_SUBTITLE = "—— 写给零基础工程师的调度算法入门（基于真实源码）"
CJK_FONT = "微软雅黑"
CODE_FONT = "Consolas"


# ---------------------------------------------------------------- 样式

def _set_cjk(style_or_run, name=CJK_FONT):
    rpr = style_or_run.font.element.get_or_add_rPr() if hasattr(style_or_run.font, "element") else None
    fonts = style_or_run.element.rPr.rFonts if hasattr(style_or_run, "element") else None
    style_or_run.font.name = name
    el = style_or_run.element if hasattr(style_or_run, "element") else style_or_run.font.element
    rPr = el.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _shade_paragraph(par, fill="F2F2F2"):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    _set_cjk(normal)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3

    for name, size, color in (("Heading 1", 18, RGBColor(0x1F, 0x3B, 0x63)),
                              ("Heading 2", 14, RGBColor(0x1F, 0x3B, 0x63)),
                              ("Heading 3", 12, RGBColor(0x33, 0x4F, 0x7C))):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        _set_cjk(st)
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)  # paragraph style
    code.base_style = doc.styles["Normal"]
    code.font.name = CODE_FONT
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    quote = doc.styles["Intense Quote"]
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
    _set_cjk(quote)


# ------------------------------------------------------- 行内 **粗体** / `代码`

_TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_runs(par, text: str) -> None:
    for part in _TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = par.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), CODE_FONT)
            rFonts.set(qn("w:hAnsi"), CODE_FONT)
        else:
            par.add_run(part)


# ---------------------------------------------------------------- 渲染

def render_blocks(doc: Document, blocks: list) -> None:
    for block in blocks:
        kind = block[0]
        if kind in ("h1", "h2", "h3"):
            doc.add_heading(block[1], level=int(kind[1]))
        elif kind == "p":
            _add_runs(doc.add_paragraph(), block[1])
        elif kind == "quote":
            par = doc.add_paragraph(style="Intense Quote")
            _add_runs(par, block[1])
        elif kind == "code":
            lines = block[1].split("\n")
            for i, line in enumerate(lines):
                par = doc.add_paragraph(style="CodeBlock")
                par.add_run(line if line else " ")
                _shade_paragraph(par)
            doc.add_paragraph(style="CodeBlock").paragraph_format.space_after = Pt(4)
        elif kind == "formula":
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(block[1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10.5)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), CODE_FONT)
            rFonts.set(qn("w:hAnsi"), CODE_FONT)
        elif kind == "table":
            spec = block[1]
            headers, rows = spec["headers"], spec["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = doc.styles["Table Grid"]
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(str(h))
                run.bold = True
                _shade_cell(cell, "D9E2F3")
            for i, row in enumerate(rows, start=1):
                for j, val in enumerate(row):
                    cell = table.rows[i].cells[j]
                    cell.paragraphs[0].text = ""
                    _add_runs(cell.paragraphs[0], str(val))
            caption = spec.get("caption")
            if caption:
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = par.add_run(caption)
                run.italic = True
                run.font.size = Pt(9)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "bullets":
            for item in block[1]:
                _add_runs(doc.add_paragraph(style="List Bullet"), item)
        elif kind == "numbers":
            for item in block[1]:
                _add_runs(doc.add_paragraph(style="List Number"), item)
        elif kind == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(f"unknown block kind: {kind}")


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# ---------------------------------------------------------------- 骨架

def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(DOC_TITLE)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(DOC_SUBTITLE)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(10):
        doc.add_paragraph()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(f"LLM4DRD 智能调度/排产平台 · 内部学习资料\n{datetime.date.today():%Y 年 %m 月 %d 日}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    par = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    hint = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "（在 Word 中右键此处 → 更新域，即可生成目录）"
    hint.append(text)
    fld.append(hint)
    par._p.append(fld)
    doc.add_page_break()


def build(output_path: str, chapter_modules: list[str]) -> None:
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    add_cover(doc)
    add_toc(doc)

    for mod_name in chapter_modules:
        mod = importlib.import_module(mod_name)
        importlib.reload(mod)
        render_blocks(doc, mod.BLOCKS)
        doc.add_page_break()

    doc.save(output_path)
    print(f"saved: {output_path}")


if __name__ == "__main__":
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    mods = sys.argv[1:] or ["chapters.ch0"]
    build("LLM4DRD排产算法基础理论.docx", mods)
